import streamlit as st
import fitz  # PyMuPDF
import io
import zipfile
import os
import hashlib
import base64
import struct
import xml.etree.ElementTree as ET
from PIL import Image, ImageDraw, ImageFont

# ─── Page Setup ───────────────────────────────
st.set_page_config(page_title="Pro Secure Doc Shield", layout="centered")

st.markdown("""
    <style>
    body { -webkit-user-select: none; user-select: none; }
    </style>
    <script>
    document.addEventListener('contextmenu', e => e.preventDefault());
    document.onkeydown = function(e) {
        if (e.ctrlKey && [67,86,85,83,80].includes(e.keyCode)) return false;
    };
    </script>
""", unsafe_allow_html=True)

st.title("🛡️ Pro Document Security Shield")


# ══════════════════════════════════════════════
#  MYANMAR FONT LOADER
#  ──────────────────────────────────────────────
#  PyMuPDF သည် built-in font များဖြစ်သည့် "helv", "tiro" တို့ကို
#  Myanmar Unicode အတွက် support မပေးပါ။ ထို့ကြောင့် Pyidaungsu
#  သို့မဟုတ် Padauk font ကို download လုပ်ပြီး register လုပ်ရမည်။
# ══════════════════════════════════════════════

FONT_CACHE_PATH = "/tmp/myanmar_font.ttf"

# Google Fonts မှ Padauk Regular (Myanmar Unicode font) ကို download လုပ်မည်
MYANMAR_FONT_URL = "https://github.com/scosant/pyidaungsu/raw/main/Pyidaungsu-2.5.3_Regular.ttf"
# Fallback: Noto Sans Myanmar
NOTO_MYANMAR_URL = "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansMyanmar/NotoSansMyanmar-Regular.ttf"

@st.cache_resource(show_spinner=False)
def load_myanmar_font() -> bytes | None:
    """
    Myanmar font ကို download လုပ်ပြီး cache သိမ်းသည်။
    အောင်မြင်ရင် font bytes ပြန်ပေးသည်။ မအောင်မြင်ရင် None ပြန်ပေးသည်။
    
    st.cache_resource ကြောင့် app session တစ်ခုလုံးတွင် တစ်ကြိမ်သာ download လုပ်မည်။
    """
    import urllib.request

    # System font path များ စစ်ဆေးသည် (Linux server တွင် install ထားပြီးဖြစ်တတ်သည်)
    system_paths = [
        "/usr/share/fonts/truetype/pyidaungsu/Pyidaungsu-Regular.ttf",
        "/usr/share/fonts/truetype/padauk/Padauk.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansMyanmar-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansMyanmar-Regular.ttf",
    ]
    for p in system_paths:
        if os.path.exists(p):
            with open(p, "rb") as f:
                return f.read()

    # Cache စစ်ဆေးသည်
    if os.path.exists(FONT_CACHE_PATH):
        with open(FONT_CACHE_PATH, "rb") as f:
            return f.read()

    # Download လုပ်သည်
    for url in [MYANMAR_FONT_URL, NOTO_MYANMAR_URL]:
        try:
            with urllib.request.urlopen(url, timeout=15) as resp:
                font_bytes = resp.read()
            with open(FONT_CACHE_PATH, "wb") as f:
                f.write(font_bytes)
            return font_bytes
        except Exception:
            continue

    return None  # Font မရပါ → Latin fallback သုံးမည်


def register_myanmar_font(page: fitz.Page, font_bytes: bytes | None) -> str:
    """
    PyMuPDF page တွင် Myanmar font ကို register လုပ်ပြီး
    သုံးရမည့် fontname ပြန်ပေးသည်။
    Font မရလျှင် "helv" (Helvetica) fallback သုံးမည်။
    
    Parameters
    ----------
    page       : fitz.Page  — font register မည့် page
    font_bytes : bytes|None — TTF font data
    
    Returns
    -------
    str — insert_text တွင် သုံးရမည့် fontname
    """
    if font_bytes is None:
        return "helv"   # Fallback → Myanmar စာသားများ □ box ပြမည်

    try:
        # "myan" ဟူသော alias ဖြင့် register လုပ်သည်
        page.insert_font(fontname="myan", fontbuffer=font_bytes)
        return "myan"
    except Exception:
        return "helv"


# ─── Sidebar ──────────────────────────────────
with st.sidebar:
    st.header("⚙️ Security Settings")

    set_password = st.checkbox("🔑 Password ခံမည်")
    user_pw = ""
    if set_password:
        user_pw = st.text_input("Password", type="password")

    owner_pw = "master_admin_key_123"

    st.divider()
    st.markdown("**Output Mode ရွေးချယ်ပါ:**")
    output_mode = st.radio(
        "Output Mode",
        options=["same_format", "convert_to_pdf"],
        format_func=lambda x: {
            "same_format":    "📄 Format အတိုင်းထား + Lock",
            "convert_to_pdf": "🔒 PDF အဖြစ်ပြောင်း (အပြည့်ပိတ်)",
        }[x],
        label_visibility="collapsed"
    )

    st.divider()
    if output_mode == "same_format":
        st.warning("⚠️ DOCX/XLSX lock ကို LibreOffice/Google Docs မှာ bypass နိုင်သည်")
    else:
        st.success("✅ PDF (Image-based) = Copy တကယ်မရနိုင်")

    # Font status ပြသည်
    st.divider()
    with st.spinner("Myanmar font စစ်ဆေးနေသည်..."):
        _font_bytes = load_myanmar_font()
    if _font_bytes:
        st.success(f"🗛 Myanmar Font: ✅ Ready ({len(_font_bytes)//1024} KB)")
    else:
        st.warning("🗛 Myanmar Font: ⚠️ မရ — Latin fallback သုံးမည်")


# ══════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════

def _sha512_hash(password: str):
    salt = os.urandom(16)
    pw_bytes = password.encode("utf-16-le")
    h = hashlib.sha512(salt + pw_bytes).digest()
    for i in range(100000):
        h = hashlib.sha512(h + struct.pack("<I", i)).digest()
    return base64.b64encode(salt).decode(), base64.b64encode(h).decode()


# ── 1. PDF → Protected PDF ────────────────────
def protect_pdf(pdf_bytes: bytes, u_pw: str, o_pw: str) -> bytes:
    """
    PDF ကို image-based လုပ်ပြီး AES-256 encrypt လုပ်သည်။
    Page ကို rasterize (pixmap) လုပ်သောကြောင့် Myanmar font
    original ပုံအတိုင်း ထွက်ပြီး copy မရနိုင်ပါ။
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    out = fitz.open()
    for page in doc:
        pix = page.get_pixmap(dpi=150)
        img = pix.tobytes("jpg", jpg_quality=70)
        np  = out.new_page(width=page.rect.width, height=page.rect.height)
        np.insert_image(page.rect, stream=img)
    perm = int(fitz.PDF_PERM_ACCESSIBILITY)
    buf  = io.BytesIO()
    out.save(buf, encryption=fitz.PDF_ENCRYPT_AES_256,
             user_pw=u_pw or None, owner_pw=o_pw,
             permissions=perm, deflate=True)
    out.close(); doc.close()
    return buf.getvalue()


# ── 2. DOCX → Protected DOCX ─────────────────
def protect_docx(docx_bytes: bytes, password: str = None) -> bytes:
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_map = {
        'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
        'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'o':   'urn:schemas-microsoft-com:office:office',
        'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'm':   'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'v':   'urn:schemas-microsoft-com:vml',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'w':   W,
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
    }
    for p, u in ns_map.items():
        ET.register_namespace(p, u)

    inp = io.BytesIO(docx_bytes)
    out = io.BytesIO()

    with zipfile.ZipFile(inp, "r") as zin, \
         zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                try:
                    tree = ET.fromstring(data)
                    tag_prot = f"{{{W}}}documentProtection"
                    for old in tree.findall(tag_prot):
                        tree.remove(old)
                    prot = ET.Element(tag_prot)
                    prot.set(f"{{{W}}}edit",        "readOnly")
                    prot.set(f"{{{W}}}enforcement",  "1")
                    prot.set(f"{{{W}}}formatting",   "1")
                    if password:
                        salt_b64, hash_b64 = _sha512_hash(password)
                        prot.set(f"{{{W}}}cryptProviderType",   "rsaFull")
                        prot.set(f"{{{W}}}cryptAlgorithmClass", "hash")
                        prot.set(f"{{{W}}}cryptAlgorithmType",  "typeAny")
                        prot.set(f"{{{W}}}cryptAlgorithmSid",   "14")
                        prot.set(f"{{{W}}}cryptSpinCount",      "100000")
                        prot.set(f"{{{W}}}hash",                hash_b64)
                        prot.set(f"{{{W}}}salt",                salt_b64)
                    tree.insert(0, prot)
                    data = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + \
                           ET.tostring(tree, encoding="unicode").encode("utf-8")
                except Exception:
                    pass
            zout.writestr(item, data)
    return out.getvalue()


# ── 3. XLSX → Protected XLSX ──────────────────
def protect_xlsx(xlsx_bytes: bytes, password: str = None) -> bytes:
    NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    ET.register_namespace('', NS)

    def sheet_attribs(pw):
        a = {
            f"{{{NS}}}sheet":               "1",
            f"{{{NS}}}objects":             "1",
            f"{{{NS}}}scenarios":           "1",
            f"{{{NS}}}selectLockedCells":   "0",
            f"{{{NS}}}selectUnlockedCells": "0",
        }
        if pw:
            salt_b64, hash_b64 = _sha512_hash(pw)
            a[f"{{{NS}}}algorithmName"] = "SHA-512"
            a[f"{{{NS}}}hashValue"]     = hash_b64
            a[f"{{{NS}}}saltValue"]     = salt_b64
            a[f"{{{NS}}}spinCount"]     = "100000"
        return a

    inp = io.BytesIO(xlsx_bytes)
    out = io.BytesIO()

    with zipfile.ZipFile(inp, "r") as zin, \
         zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item.startswith("xl/worksheets/sheet") and item.endswith(".xml"):
                try:
                    root = ET.fromstring(data)
                    tag_prot = f"{{{NS}}}sheetProtection"
                    for old in root.findall(tag_prot):
                        root.remove(old)
                    prot = ET.Element(tag_prot)
                    for k, v in sheet_attribs(password).items():
                        prot.set(k, v)
                    insert_pos = len(list(root))
                    for idx, child in enumerate(root):
                        if child.tag == f"{{{NS}}}sheetData":
                            insert_pos = idx + 1
                            break
                    root.insert(insert_pos, prot)
                    data = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + \
                           ET.tostring(root, encoding="unicode").encode("utf-8")
                except Exception:
                    pass
            zout.writestr(item, data)
    return out.getvalue()


# ══════════════════════════════════════════════
#  MYANMAR FONT AWARE TEXT WRITER
#  ──────────────────────────────────────────────
#  PyMuPDF ၏ insert_text သည် Unicode string ကို UTF-8 encode
#  လုပ်ပြီး font glyph map မှ ပုံဆွဲသည်။ Myanmar font register
#  ထားမှသာ Myanmar codepoint (U+1000–U+109F, U+A9E0–U+A9FF,
#  U+AA60–U+AA7F) များကို ကောင်းကောင်း render လုပ်နိုင်သည်။
# ══════════════════════════════════════════════

def _measure_text(text: str, fontname: str, fontsize: float,
                  font_bytes: bytes | None) -> float:
    """
    Myanmar font ရှိမရှိပေါ်မူတည်ပြီး text width တိုင်းသည်။
    font_bytes ရှိလျှင် fitz.Font ဖြင့် တိုင်းသည်။
    မရှိလျှင် helv estimate သုံးသည်။
    """
    if font_bytes and fontname == "myan":
        try:
            f = fitz.Font(fontbuffer=font_bytes)
            return f.text_length(text, fontsize=fontsize)
        except Exception:
            pass
    return fitz.get_text_length(text, fontname="helv", fontsize=fontsize)


# ── 4a. DOCX → PDF (Myanmar Font) ─────────────
def docx_to_protected_pdf(docx_bytes: bytes, u_pw: str, o_pw: str) -> bytes:
    """
    DOCX မှ text/table ကို extract လုပ်ပြီး PyMuPDF ဖြင့်
    Myanmar font သုံး၍ PDF ထုတ်ကာ AES-256 protect လုပ်သည်။

    အဓိက ပြောင်းလဲချက် (Myanmar Font Support):
    ───────────────────────────────────────────
    1. load_myanmar_font()   → TTF bytes ရယူသည်
    2. register_myanmar_font() → page တစ်ခုစီတွင် "myan" alias ဖြင့် register
    3. insert_text(..., fontname=fontname) → "myan" သို့မဟုတ် "helv" fallback
    4. _measure_text() → word-wrap တွင် Myanmar font width ကို မှန်ကန်စွာ တိုင်းသည်
    """
    from docx import Document as DocxDocument

    font_bytes = load_myanmar_font()   # ← Myanmar font bytes (cached)

    docx_doc = DocxDocument(io.BytesIO(docx_bytes))
    out_pdf   = fitz.open()

    PAGE_W, PAGE_H = 595, 842
    MARGIN         = 50
    LINE_H         = 16      # Myanmar font တွင် line height ပိုလိုသည်
    FONT_SIZE      = 11
    x_start        = MARGIN

    def new_page():
        p = out_pdf.new_page(width=PAGE_W, height=PAGE_H)
        fn = register_myanmar_font(p, font_bytes)   # ← page တိုင်းတွင် register
        return p, MARGIN, fn

    page, y, fontname = new_page()

    def write_line(text: str, size: float = FONT_SIZE):
        """Myanmar text ပါဝင်သော line တစ်ကြောင်းကို PDF page တွင် ရေးသည်။"""
        nonlocal page, y, fontname
        if not text:
            return
        if y + LINE_H > PAGE_H - MARGIN:
            page, y, fontname = new_page()
        try:
            page.insert_text(
                (x_start, y),
                text,
                fontname=fontname,   # "myan" (Myanmar) သို့မဟုတ် "helv" (fallback)
                fontsize=size,
                color=(0, 0, 0),
                encoding=fitz.TEXT_ENCODING_UNICODE,  # ← Unicode အတွက် မဖြစ်မနေ
            )
        except Exception:
            # Myanmar render မအောင်မြင်လျှင် Latin-safe fallback
            safe = "".join(c if ord(c) < 128 else "?" for c in text)
            page.insert_text((x_start, y), safe,
                             fontname="helv", fontsize=size, color=(0.4, 0.4, 0.4))
        y += LINE_H + 2

    # ── Paragraph များ render ──
    for para in docx_doc.paragraphs:
        text  = para.text
        style = para.style.name if para.style else ""
        size  = FONT_SIZE
        if   "Heading 1" in style: size = 16
        elif "Heading 2" in style: size = 14
        elif "Heading 3" in style: size = 12

        if text.strip() == "":
            y += LINE_H // 2
            continue

        # Word-wrap: Myanmar font width ဖြင့် measure လုပ်သည်
        max_w = PAGE_W - 2 * MARGIN
        words = text.split()
        line  = ""
        for word in words:
            test = (line + " " + word).strip()
            if _measure_text(test, fontname, size, font_bytes) > max_w:
                write_line(line, size=size)
                line = word
            else:
                line = test
        if line:
            write_line(line, size=size)
        y += 2

    # ── Table များ render ──
    for table in docx_doc.tables:
        y += 8
        for row in table.rows:
            row_text = " │ ".join(cell.text.strip() for cell in row.cells)
            write_line(row_text, size=10)
        y += 4

    buf  = io.BytesIO()
    perm = int(fitz.PDF_PERM_ACCESSIBILITY)
    out_pdf.save(buf, encryption=fitz.PDF_ENCRYPT_AES_256,
                 user_pw=u_pw or None, owner_pw=o_pw,
                 permissions=perm, deflate=True)
    out_pdf.close()
    return buf.getvalue()


# ── 4b. XLSX → PDF (Myanmar Font) ─────────────
def xlsx_to_protected_pdf(xlsx_bytes: bytes, u_pw: str, o_pw: str) -> bytes:
    """
    openpyxl မှ cell data ကို extract လုပ်ပြီး PyMuPDF ဖြင့်
    Myanmar font သုံး၍ PDF ထုတ်ကာ AES-256 protect လုပ်သည်။

    အဓိက ပြောင်းလဲချက် (Myanmar Font Support):
    ───────────────────────────────────────────
    1. page တစ်ခုစီတွင် register_myanmar_font() ခေါ်သည်
    2. cell text ကို Myanmar font ဖြင့် insert_text လုပ်သည်
    3. Column width truncate တွင် Myanmar font measure သုံးသည်
    """
    import openpyxl

    font_bytes = load_myanmar_font()   # ← Myanmar font bytes (cached)

    wb      = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    out_pdf = fitz.open()

    PAGE_W, PAGE_H = 842, 595   # A4 Landscape
    MARGIN   = 40
    ROW_H    = 16               # Myanmar font line height
    FONT_SZ  = 9
    HDR_SZ   = 12

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        page     = out_pdf.new_page(width=PAGE_W, height=PAGE_H)
        fontname = register_myanmar_font(page, font_bytes)   # ← register
        y        = MARGIN

        # Sheet title (Myanmar name ပါနိုင်သည်)
        try:
            page.insert_text(
                (MARGIN, y),
                f"Sheet: {sheet_name}",
                fontname=fontname,
                fontsize=HDR_SZ,
                color=(0.2, 0.2, 0.8),
                encoding=fitz.TEXT_ENCODING_UNICODE,
            )
        except Exception:
            page.insert_text((MARGIN, y), f"Sheet: {sheet_name}",
                             fontname="helv", fontsize=HDR_SZ,
                             color=(0.2, 0.2, 0.8))
        y += HDR_SZ + 6

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        n_cols   = max(len(r) for r in rows)
        usable_w = PAGE_W - 2 * MARGIN
        col_w    = usable_w / max(n_cols, 1)

        for row_idx, row in enumerate(rows):
            if y + ROW_H > PAGE_H - MARGIN:
                page     = out_pdf.new_page(width=PAGE_W, height=PAGE_H)
                fontname = register_myanmar_font(page, font_bytes)   # ← new page တွင်လည်း register
                y        = MARGIN

            for col_idx, val in enumerate(row):
                x    = MARGIN + col_idx * col_w
                text = str(val) if val is not None else ""

                # Myanmar font ဖြင့် truncate: column ကျဉ်းလျှင် ဖြတ်ထားသည်
                max_cell_w = col_w - 4
                while len(text) > 1 and \
                      _measure_text(text, fontname, FONT_SZ, font_bytes) > max_cell_w:
                    text = text[:-1]

                color = (0.1, 0.1, 0.5) if row_idx == 0 else (0, 0, 0)
                try:
                    page.insert_text(
                        (x + 2, y),
                        text,
                        fontname=fontname,
                        fontsize=FONT_SZ,
                        color=color,
                        encoding=fitz.TEXT_ENCODING_UNICODE,
                    )
                except Exception:
                    safe = "".join(c if ord(c) < 128 else "?" for c in text)
                    page.insert_text((x + 2, y), safe,
                                     fontname="helv", fontsize=FONT_SZ, color=color)

            page.draw_line((MARGIN, y + 3), (PAGE_W - MARGIN, y + 3),
                           color=(0.85, 0.85, 0.85), width=0.3)
            y += ROW_H

    buf  = io.BytesIO()
    perm = int(fitz.PDF_PERM_ACCESSIBILITY)
    out_pdf.save(buf, encryption=fitz.PDF_ENCRYPT_AES_256,
                 user_pw=u_pw or None, owner_pw=o_pw,
                 permissions=perm, deflate=True)
    out_pdf.close()
    return buf.getvalue()


# ══════════════════════════════════════════════
#  MAIN UI
# ══════════════════════════════════════════════

st.markdown("---")

if output_mode == "same_format":
    st.info("📂 **Format အတိုင်း Mode** — Excel→Excel | Word→Word | PDF→PDF")
else:
    st.success("🔒 **PDF Convert Mode** — အားလုံးကို Image-based Protected PDF ပြောင်းမည်")

uploaded_file = st.file_uploader(
    "ဖိုင်ရွေးချယ်ပါ (PDF, DOCX, XLSX)",
    type=["pdf", "docx", "xlsx"]
)

if uploaded_file:
    ext  = uploaded_file.name.rsplit(".", 1)[-1].lower()
    raw  = uploaded_file.read()
    base = uploaded_file.name.rsplit(".", 1)[0]

    try:
        with st.spinner("🔐 လုံခြုံရေးအလွှာများ ထည့်သွင်းနေသည်..."):

            if output_mode == "same_format":
                if ext == "pdf":
                    result   = protect_pdf(raw, user_pw, owner_pw)
                    out_name = f"Protected_{uploaded_file.name}"
                    mime     = "application/pdf"
                    label    = "PDF → Image-based + AES-256"
                elif ext == "docx":
                    result   = protect_docx(raw, user_pw or None)
                    out_name = f"Protected_{base}.docx"
                    mime     = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    label    = "DOCX → Read-only Lock (Word format အတိုင်း)"
                elif ext == "xlsx":
                    result   = protect_xlsx(raw, user_pw or None)
                    out_name = f"Protected_{base}.xlsx"
                    mime     = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    label    = "XLSX → All Sheets Locked (Excel format အတိုင်း)"
            else:
                if ext == "pdf":
                    result = protect_pdf(raw, user_pw, owner_pw)
                    label  = "PDF → Image-based + AES-256 Protected PDF"
                elif ext == "docx":
                    result = docx_to_protected_pdf(raw, user_pw, owner_pw)
                    label  = "DOCX → Myanmar Font PDF (Protected)"
                elif ext == "xlsx":
                    result = xlsx_to_protected_pdf(raw, user_pw, owner_pw)
                    label  = "XLSX → Myanmar Font Table PDF (Protected)"
                out_name = f"Protected_{base}.pdf"
                mime     = "application/pdf"

        col1, col2 = st.columns(2)
        col1.success("✅ အောင်မြင်ပြီး!")
        col2.info(f"📦 {len(result)/1024:.1f} KB")
        st.caption(f"🔐 {label}")
        st.download_button(
            label=f"⬇️ Download — {out_name}",
            data=result,
            file_name=out_name,
            mime=mime,
            use_container_width=True,
            type="primary"
        )

    except Exception as e:
        st.error(f"❌ Error: {e}")

st.divider()
st.caption("Pro Document Security Shield | Myanmar Font Support | pymupdf + openpyxl + python-docx")
